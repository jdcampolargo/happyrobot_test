from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Acme_Logistics_Build_Description.docx"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if bold:
        run.font.color.rgb = None
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        shade_cell(hdr[idx], "E5E7EB")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Acme Logistics\nInbound Carrier Sales Automation")
    r.bold = True
    r.font.size = Pt(24)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Build description for freight brokerage evaluation")
    sr.font.size = Pt(12)
    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This proof of concept automates the first-line inbound carrier sales workflow: carrier eligibility, "
        "load matching, price negotiation, mocked transfer, and operational reporting. HappyRobot handles the "
        "live voice conversation while a secured custom API owns deterministic broker logic and metrics."
    )

    doc.add_heading("Customer Problem", level=1)
    doc.add_paragraph(
        "Inbound carrier calls are repetitive but operationally sensitive. A brokerage needs speed, consistency, "
        "and pricing discipline without allowing an AI agent to invent carrier eligibility or rate decisions."
    )
    add_bullets(doc, [
        "Carrier reps spend time collecting the same MC, lane, equipment, and offer information.",
        "Eligibility checks and rate decisions vary by rep unless encoded as policy.",
        "Post-call data is often incomplete, making conversion and lane-performance analysis weak.",
        "A useful PoC must show real workflow behavior, not just a scripted voice demo.",
    ])

    doc.add_heading("Solution Architecture", level=1)
    add_numbered(doc, [
        "Carrier starts a HappyRobot web call.",
        "Voice agent collects MC number and calls the verification tool.",
        "If eligible, the agent collects lane/equipment preferences and calls load search.",
        "The agent pitches the best matching load and asks for acceptance.",
        "Counteroffers are evaluated by the pricing policy API for up to three rounds.",
        "Accepted loads use the required mocked transfer phrase.",
        "Post-call AI extraction and classification create a structured call record.",
        "The custom dashboard reads logged call records from the backend database.",
    ])

    add_table(doc, ["Layer", "Responsibility", "Implementation"], [
        ["HappyRobot workflow", "Live voice call, tool orchestration, post-call extraction/classification", "Web Call trigger, voice prompt node, tool nodes, AI Extract, AI Classify, HTTP action"],
        ["Custom API", "Carrier verification, load search, negotiation policy, call logging", "FastAPI service with API-key authentication"],
        ["Data", "Load inventory and call metrics", "Local JSON load file plus SQLite call log for the PoC"],
        ["Dashboard", "Broker-specific operational reporting", "HTML dashboard backed by /api/metrics"],
    ])

    doc.add_heading("Backend Capabilities", level=1)
    add_table(doc, ["Capability", "Endpoint", "Purpose"], [
        ["Carrier verification", "POST /api/carriers/verify", "Checks MC/docket number, operating authority, out-of-service state, and eligibility reason."],
        ["Load search", "POST /api/loads/search", "Finds viable loads by origin, destination, pickup date, and equipment type."],
        ["Negotiation evaluation", "POST /api/offers/evaluate", "Accepts, counters, or rejects carrier rates using deterministic broker policy."],
        ["Call logging", "POST /api/calls/log", "Stores extracted offer data, outcome, sentiment, transcript, and duration."],
        ["Metrics", "GET /api/metrics", "Returns aggregate and call-level data for the custom dashboard."],
    ])

    doc.add_heading("Negotiation Policy", level=1)
    doc.add_paragraph(
        "The model does not decide prices on its own. It must call the negotiation endpoint whenever the carrier "
        "offers a rate above the listed loadboard rate. The API computes the maximum acceptable rate and returns "
        "the exact message the agent should use."
    )
    add_bullets(doc, [
        "Default target is the loadboard rate.",
        "Maximum acceptable rate is 108% for high-priority or urgent freight, 105% for near-pickup freight, and 103% for standard freight.",
        "Counters are rounded to the nearest $25.",
        "The agent handles no more than three negotiation back-and-forths.",
        "After round three, above-policy offers are rejected or routed for human follow-up.",
    ])

    doc.add_heading("Dashboard and Metrics", level=1)
    doc.add_paragraph(
        "The dashboard is intentionally independent from HappyRobot analytics. It is a broker-facing operating view "
        "of the use case, not just a generic call analytics page."
    )
    add_bullets(doc, [
        "Total inbound calls and booked calls.",
        "Mocked transfer rate and qualified-carrier rate.",
        "Average final offer and average rate versus loadboard.",
        "Outcome distribution and sentiment distribution.",
        "Lane-level calls, bookings, conversion rate, and average final offer.",
        "Recent call-level records with MC number, load, outcome, sentiment, offers, and summary.",
    ])

    doc.add_heading("Security and Deployment", level=1)
    add_bullets(doc, [
        "Every API endpoint requires an API key via X-API-Key, bearer token, or dashboard query parameter for demo access.",
        "The API is containerized with Docker and can be deployed to Railway, Render, Fly.io, AWS, Azure, or Google Cloud.",
        "HTTPS can be provided by the cloud platform or by the included Caddy reverse proxy for self-hosted/local demos.",
        "Secrets are environment variables and are not committed to the repository.",
        "Persistent storage can be mounted at /app/data to preserve SQLite call logs across deploys.",
    ])

    doc.add_heading("Demo Plan", level=1)
    doc.add_heading("Successful booking", level=2)
    add_bullets(doc, [
        "Carrier gives MC 123456.",
        "Agent verifies eligibility.",
        "Carrier asks for Chicago, IL to Dallas, TX dry van freight.",
        "Agent pitches load ACME-1001 at $2,450.",
        "Carrier counters above the listed rate.",
        "Agent calls the negotiation endpoint and counters within policy.",
        "Carrier accepts an acceptable rate.",
        "Agent uses the mocked transfer phrase and the dashboard logs booked_transfer_mocked.",
    ])

    doc.add_heading("Ineligible carrier", level=2)
    add_bullets(doc, [
        "Carrier gives MC 654321.",
        "Verification fails eligibility.",
        "Agent refuses to quote freight automatically.",
        "Post-call classification logs carrier_ineligible for reporting.",
    ])

    doc.add_heading("Production Roadmap", level=1)
    add_bullets(doc, [
        "Replace the local load file with the broker's TMS or loadboard source.",
        "Add insurance, fraud, packet-completion, lane-history, and facility-access checks.",
        "Replace demo API-key dashboard access with SSO and role-based access control.",
        "Write accepted tenders back to the TMS and create sales-rep tasks.",
        "Add real transfer behavior after phone-number and telephony integration are enabled.",
        "Add QA sampling, alerting, and model/policy performance review workflows.",
    ])

    doc.add_heading("Success Criteria", level=1)
    add_bullets(doc, [
        "The web call agent gathers MC number before discussing freight.",
        "Carrier eligibility controls whether freight can be quoted.",
        "Structured load data is searched and pitched accurately.",
        "Counteroffers are constrained to three rounds and evaluated by the API.",
        "Accepted rates trigger the required mocked transfer phrase.",
        "Post-call extraction, outcome classification, and sentiment classification are logged.",
        "A custom dashboard reports broker-specific metrics outside platform analytics.",
    ])

    doc.core_properties.title = "Acme Logistics Inbound Carrier Sales Automation"
    doc.core_properties.subject = "HappyRobot FDE technical challenge build description"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
